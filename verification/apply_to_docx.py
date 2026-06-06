# -*- coding: utf-8 -*-
"""
Aplică conținutul nou (new_content.SECTIONS) în documentul Word, păstrând stilul.

Pentru fiecare heading-ancoră:
  1. Găsește paragraful-heading (după text exact).
  2. Parcurge paragrafele de CONȚINUT până la următorul heading.
  3. Șterge DOAR paragrafele text (păstrează cele cu imagini).
  4. Inserează paragrafele noi imediat după heading, cu stil "Body Text",
     Times New Roman 12pt, line spacing 1.5, justified.
  Tabelele (obiecte separate de paragrafe) rămân neatinse pe poziția lor.

Heading-urile NOI (h4) din conținut se inserează cu stil bold 12pt.
"""
import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from new_content import SECTIONS

DOCX = r"D:/OneDrive - Realworld Holding b.v/Documents/67/Word/Paun_Raluca_Raport_CS4-final.docx"

doc = Document(DOCX)
body = doc.element.body


def has_image(p):
    return len(p._p.findall(".//" + qn("a:blip"))) > 0


# Heading-uri H4 ne-numerotate care trebuie tratate ca opritori (delimitează secțiuni).
# Acestea sunt și ancore în SECTIONS sau sub-titluri ale documentului.
_H4_STOPPERS = {
    "achiziția datelor:",
    "dataset principal: purtători de proteză transtibială – samala et al. (2024)",
    "dataset complementar: date imu în condiții reale de utilizare - wassall ntnu (2025)",
    "gaitrec",
    "caracteristici principale:",
    "modul de preprocesare imu",
    "modul de detecție a fazelor ciclului de mers",
    "modul de extragere a parametrilor de mers",
    "modul de control al gleznei bazat pe faze (fsm)",
    "modul de vizualizare și raportare (dashboard)",
}


def is_heading_text(t):
    """Heuristică: e un heading numerotat (4.x / 5.x) sau un titlu de secțiune scurt."""
    import re
    t = t.strip()
    if re.match(r"^\d+\.\d+", t):
        return True
    if t.lower() in _H4_STOPPERS:
        return True
    # titluri majuscule scurte de secțiune
    if t and t.isupper() and len(t) < 80 and t not in ("HS", "TO", "FSM", "IMU", "OMC"):
        return True
    return False


def find_para_index(paras, text):
    for i, p in enumerate(paras):
        if p.text.strip() == text.strip():
            return i
    # fallback: startswith pe primele 40 caractere
    key = text.strip()[:40]
    for i, p in enumerate(paras):
        if p.text.strip().startswith(key):
            return i
    return -1


def style_body(p):
    p.style = doc.styles["Body Text"]
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        # asigură font pentru caractere est-europene
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def style_h4(p):
    p.style = doc.styles["Normal"]
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.bold = True


def insert_paragraph_after(paragraph, text, kind):
    """Inserează un paragraf nou imediat DUPĂ `paragraph`. Returnează noul paragraf."""
    new_p = copy.deepcopy(paragraph._p)
    # golește conținutul copiat
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np_ = Paragraph(new_p, paragraph._parent)
    run = np_.add_run(text)
    if kind == "h4":
        style_h4(np_)
    else:
        style_body(np_)
    return np_


total_inserted = 0
total_deleted = 0
report = []

for anchor, blocks in SECTIONS.items():
    paras = doc.paragraphs
    idx = find_para_index(paras, anchor)
    if idx < 0:
        report.append(f"  ⚠ ANCORĂ NEGĂSITĂ: {anchor[:50]}")
        continue

    heading_p = paras[idx]

    # 1) Identifică și șterge paragrafele de conținut text până la următorul heading
    j = idx + 1
    deleted_here = 0
    while j < len(paras):
        p = paras[j]
        t = p.text.strip()
        if is_heading_text(t):
            break  # am ajuns la următorul heading
        if has_image(p):
            j += 1
            continue  # PĂSTREAZĂ paragrafele cu imagini
        if t == "":
            j += 1
            continue  # păstrează paragrafele goale (spațiere)
        # șterge paragraful text
        p._p.getparent().remove(p._p)
        deleted_here += 1
        paras = doc.paragraphs  # reîmprospătează lista după ștergere
        # nu incrementăm j: lista s-a scurtat, dar recalculăm idx-ul heading-ului
        idx = find_para_index(paras, anchor)
        heading_p = paras[idx]
        j = idx + 1
        # re-scan de la început pentru siguranță (volum mic)
    total_deleted += deleted_here

    # 2) Inserează noile paragrafe imediat după heading, în ordine inversă
    #    (fiecare se inserează DUPĂ heading, deci ultimul scris ajunge primul →
    #     inserăm în ordine inversă ca să iasă în ordinea corectă)
    anchor_p = heading_p
    for kind, text in blocks:
        anchor_p = insert_paragraph_after(anchor_p, text, kind)
        total_inserted += 1

    report.append(f"  ✓ {anchor[:48]:48s} | șters {deleted_here}, inserat {len(blocks)}")

doc.save(DOCX)
print("APLICAT.")
print(f"Total paragrafe șterse: {total_deleted}, inserate: {total_inserted}")
for r in report:
    print(r)
